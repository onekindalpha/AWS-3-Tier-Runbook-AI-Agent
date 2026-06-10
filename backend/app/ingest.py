from __future__ import annotations

from dataclasses import dataclass, asdict
from pathlib import Path
from docx import Document
import json
import re

from .config import DOCX_PATH, INDEX_PATH, DATA_DIR
from .text_utils import normalize, clean_paragraph, is_noise_paragraph


PAGE_KEYWORDS = [
    ("network", ["vpc", "subnet", "route table", "internet gateway", "nat gateway", "cidr", "network", "서브넷", "라우팅", "게이트웨이"]),
    ("security", ["security group", "nacl", "inbound", "outbound", "port", "ssh", "http", "https", "보안", "인바운드", "아웃바운드"]),
    ("web-api-db", ["web tier", "web", "nginx", "public alb", "load balancer", "target group", "api", "flask", "rds", "database", "mysql", "웹", "로드밸런서", "백엔드", "데이터베이스"]),
    ("validation", ["validation", "verify", "checklist", "health", "status", "test", "curl", "검증", "테스트", "상태"]),
    ("overview", ["overview", "architecture", "3-tier", "3 tier", "아키텍처", "전체", "구조"]),
]


@dataclass
class Chunk:
    chunk_id: str
    page: str
    section: str
    content: str


def detect_page(text: str) -> str:
    lower = text.lower()
    best_page = "overview"
    best_score = 0
    for page, keywords in PAGE_KEYWORDS:
        score = 0
        for kw in keywords:
            if kw.lower() in lower:
                score += 4 if " " in kw else 1
        if score > best_score:
            best_page = page
            best_score = score
    return best_page


def detect_section(text: str, last_section: str) -> str:
    lower = text.lower()
    if any(k in lower for k in ["vpc", "subnet", "route table", "internet gateway", "nat"]):
        return "Network 설정"
    if any(k in lower for k in ["security group", "nacl", "inbound", "outbound", "port"]):
        return "Security 규칙"
    if any(k in lower for k in ["alb", "load balancer", "target group", "nginx", "web tier"]):
        return "Web Tier 연결"
    if any(k in lower for k in ["api", "flask", "application tier", "backend"]):
        return "API Tier 연결"
    if any(k in lower for k in ["rds", "database", "mysql", "db subnet"]):
        return "Database 연결"
    if any(k in lower for k in ["curl", "health", "status", "test", "validation", "verify"]):
        return "검증 절차"
    if 8 <= len(text) <= 90 and not text.endswith(".") and not text.endswith("다."):
        return text[:80]
    return last_section or "매뉴얼 본문"


def load_docx_paragraphs(docx_path: Path = DOCX_PATH) -> list[str]:
    document = Document(docx_path)
    paragraphs: list[str] = []

    for para in document.paragraphs:
        text = clean_paragraph(para.text)
        if not is_noise_paragraph(text):
            paragraphs.append(text)

    for table in document.tables:
        for row in table.rows:
            row_text = " ".join(clean_paragraph(cell.text) for cell in row.cells)
            row_text = clean_paragraph(row_text)
            if not is_noise_paragraph(row_text):
                paragraphs.append(row_text)

    return paragraphs


def build_chunks(paragraphs: list[str], max_chars: int = 850, overlap_paragraphs: int = 1) -> list[Chunk]:
    chunks: list[Chunk] = []
    buffer: list[str] = []
    last_section = "매뉴얼 본문"

    def flush() -> None:
        nonlocal buffer, last_section
        joined = normalize(" ".join(buffer))
        if len(joined) >= 45:
            chunk_id = f"chunk-{len(chunks) + 1:04d}"
            chunks.append(
                Chunk(
                    chunk_id=chunk_id,
                    page=detect_page(joined),
                    section=last_section,
                    content=joined[:1200],
                )
            )
        buffer = buffer[-overlap_paragraphs:] if overlap_paragraphs > 0 else []

    for para in paragraphs:
        last_section = detect_section(para, last_section)
        buffer.append(para)
        if len(normalize(" ".join(buffer))) >= max_chars or len(buffer) >= 4:
            flush()

    if buffer:
        flush()

    return chunks


def write_index(chunks: list[Chunk], index_path: Path = INDEX_PATH) -> None:
    DATA_DIR.mkdir(parents=True, exist_ok=True)
    payload = {
        "source": str(DOCX_PATH),
        "chunk_count": len(chunks),
        "chunks": [asdict(chunk) for chunk in chunks],
    }
    index_path.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")


def load_or_build_index(force: bool = False) -> list[dict]:
    if INDEX_PATH.exists() and not force:
        payload = json.loads(INDEX_PATH.read_text(encoding="utf-8"))
        return payload.get("chunks", [])

    paragraphs = load_docx_paragraphs()
    chunks = build_chunks(paragraphs)
    write_index(chunks)
    return [asdict(chunk) for chunk in chunks]


if __name__ == "__main__":
    chunks = load_or_build_index(force=True)
    print(f"Indexed {len(chunks)} chunks from {DOCX_PATH}")
